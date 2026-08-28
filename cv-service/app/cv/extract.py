"""Pull text out of an uploaded CV.

Why this is a tool and not just "paste the PDF into the prompt": a two-page CV
is 1.5-3k tokens of raw dump, most of it layout noise — repeated headers, page
numbers, stray glyphs from a two-column layout. Sectioning it here and capping
each section means the model sees a fraction of that, which is the whole reason
a small model can do this job.

Extraction is heuristic and says so. It never claims a field is right; it hands
the model labelled candidates and the model confirms them with the user.
"""
from __future__ import annotations

import io
import re

from . import layout
from .photo import (
    extract_page_image,
    extract_portrait,
    extract_portrait_from_docx,
    render_pdf_page,
)
from .quality import Assessment, Grade, assess

# Enough to carry a section's substance, short enough that a padded CV cannot
# blow the context window. Measured against real two-page CVs.
SECTION_CAP = 1200
TOTAL_CAP = 6000

EMAIL_RE = re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b")

# Anchored on '+' or a leading 0 rather than matching loose digit groups.
# Two reasons. A pattern of fixed-width groups drops the country code on the
# very common '+212 6 23 84 25 35' shape, because the '6' is a single digit —
# and a CV that loses its country code is worse than one with no phone at all.
# Anchoring also keeps '2022-2027' out: a date range starts with neither.
PHONE_RE = re.compile(r"\+\d[\d\s.\-()]{6,20}\d|\b0\d[\d\s.\-()]{6,18}\d")
URL_RE = re.compile(r"\b(?:https?://|www\.)\S+|\b(?:github|linkedin)\.com/\S+", re.I)

# Heading vocabulary, EN + FR — the two languages the templates support.
SECTION_PATTERNS: dict[str, tuple[str, ...]] = {
    # Without this the whole contact block of any CV that labels it — which is
    # most designed ones — was dropped, taking the city and any link that is
    # not an email, phone or URL with it.
    "contact": ("contact", "contacts", "coordonnées", "details", "info"),
    "profile": ("profile", "summary", "about", "objective", "profil", "résumé", "à propos"),
    "experience": ("experience", "employment", "work history", "professional", "expérience", "parcours"),
    "internships": ("internship", "stage", "stages"),
    "education": ("education", "academic", "qualification", "formation", "études", "diplôme"),
    "skills": ("skills", "competencies", "technical", "technologies", "compétences", "outils"),
    "languages": ("languages", "langues"),
    "certifications": ("certification", "certificate", "licences", "certificats"),
    "projects": ("projects", "portfolio", "projets", "réalisations"),
    "interests": ("interests", "hobbies", "centres d'intérêt", "loisirs"),
}


# "Programming Languages", "Languages & Frameworks", "Languages and Tools" —
# a *skills sub-heading*, not the spoken-languages section.
#
# This cost a visitor their entire skills section. Their CV's sidebar read
# "TECHNICAL SKILLS" then "LANGUAGES & FRAMEWORKS" then the technologies. The
# first line opened `skills` correctly; the second matched the word
# "languages" and closed it again immediately — so `skills` ended up empty and
# was dropped, and every technology was filed under Languages next to "Arabic
# — Native". The coherence check could not catch it either: the real
# "LANGUAGES" heading later appended the actual spoken languages to the same
# key, so the section did contain language-ish text and looked fine.
#
# Matched *before* the general vocabulary and mapped to `skills`, not merely
# rejected as a heading: a CV whose only such heading is a standalone
# "Programming Languages" (no "Skills" heading at all) then still files those
# technologies under skills rather than dropping them into whatever section
# happened to be open.
_SKILLS_SUBHEADING = re.compile(
    r"\b(?:programming|programmation|coding|scripting|markup|query|"
    r"frameworks?|tools?|technolog\w*|librar\w*|software|development)\b",
    re.I,
)

# The word that needs qualifying before it means "spoken languages".
#
# French is the useful case here: it separates the two senses lexically —
# *langage* is a programming language, *langue* is one you speak — so
# "Langages" needs no qualifier at all to be a skills heading, while "Langues"
# is unambiguously the spoken section.
_LANGUAGE_WORD = re.compile(r"(?<![\w'])(?:languages?|langues?)(?![\w'])", re.I)
_PROGRAMMING_LANGUAGE_WORD = re.compile(r"(?<![\w'])langages?(?![\w'])", re.I)


class ExtractionError(Exception):
    """Raised when the file cannot yield text at all."""


def _text_from_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency is declared
        raise ExtractionError("PDF support is not installed on the server.") from exc

    try:
        reader = PdfReader(io.BytesIO(data))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as exc:
        # The exception name means nothing to whoever uploaded the file, so the
        # message says what to do instead. The type is kept for the log only.
        raise ExtractionError(
            "That PDF appears to be damaged or password-protected, so it could "
            "not be opened. Try re-saving or re-exporting it, or type your "
            "details in the chat and I'll build the CV from those."
        ) from exc


def _text_from_docx(data: bytes) -> str:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover - dependency is declared
        raise ExtractionError("DOCX support is not installed on the server.") from exc

    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise ExtractionError(f"Could not read that DOCX ({type(exc).__name__}).") from exc

    parts = [p.text for p in document.paragraphs]
    # Many CVs lay everything out in a table; ignoring tables loses the whole CV.
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _looks_letter_spaced(chunk: str) -> bool:
    """True when a chunk reads like "R I C H A R D" — a real word or phrase
    with a space forced between every letter, rather than actual short words.

    Seen on real uploads exported from Canva-style design tools: certain text
    runs in the PDF (a tracked headline, a whole paragraph, a section title —
    inconsistently, depending on the run's embedded font) come out of pypdf's
    extractor with one space per glyph. Left alone, this defeats both the
    section-heading matcher ("P R O F I L E S U M M A R Y" contains no
    contiguous "profile") and the name heuristic (a 14-token "word" fails the
    2-5-word name check, so a later ordinary line — an address, in one real
    case — gets picked as the name instead).

    Guarded to a high single-character ratio over at least 3 tokens so this
    cannot misfire on genuinely short real words ("AI ML NLP dev").
    """
    tokens = chunk.split(" ")
    if len(tokens) < 3:
        return False
    single_char = sum(1 for token in tokens if len(token) == 1)
    return single_char / len(tokens) >= 0.7


def _repair_letter_spacing(raw_line: str) -> str:
    """Undo `_looks_letter_spaced` runs, using the one signal that still
    distinguishes "space between letters" from "space between words" at this
    point: pypdf's extractor renders the former as a single space and the
    latter as a double space (or more), because it derives spacing from the
    glyphs' actual on-page gaps. That distinction is lost the moment this
    line's whitespace gets collapsed to single spaces (`_clean`'s next step),
    so the repair has to happen first, on the still-doubled-space raw line.
    """
    chunks = re.split(r" {2,}", raw_line.strip())
    repaired = [
        chunk.replace(" ", "") if _looks_letter_spaced(chunk) else chunk
        for chunk in chunks
    ]
    return " ".join(repaired)


def _clean(text: str) -> str:
    """Drop layout noise that survives extraction."""
    lines: list[str] = []
    seen_header: dict[str, int] = {}
    for raw in text.splitlines():
        line = " ".join(_repair_letter_spacing(raw).split())
        if not line:
            continue
        # Bare page numbers, and "Page 1 of 3". Bounded to four digits: an
        # unbounded \d+ also matches a phone number written without separators
        # ("0623842535"), which would silently delete it from the CV.
        if re.fullmatch(r"(page\s*)?\d{1,4}(\s*(/|of|sur)\s*\d{1,4})?", line, re.I):
            continue
        # A short line repeated on every page is a running header or footer.
        if len(line) < 60:
            seen_header[line] = seen_header.get(line, 0) + 1
            if seen_header[line] > 2:
                continue
        lines.append(line)
    return "\n".join(lines)


def _looks_like_heading(line: str) -> str | None:
    """Return the canonical section name if this line reads as its heading.

    Two things this has to survive, both found on real CVs:

    * **Markdown.** People paste CVs written in Markdown, where the heading is
      "## Work Experience" or "**SKILLS**". The markers carry the meaning — they
      are what makes it a heading — so they are stripped, not matched against.
    * **Qualified headings.** The needle can sit anywhere in the phrase.
      "Work Experience" is probably the commonest heading in an English CV and
      it does not *begin* with "experience", so prefix matching missed it
      entirely and dropped the whole section.

    Matched on whole words, so "Experienced professional" is not a heading and
    "Skills" inside a sentence cannot open a section. Length still bounds it:
    a heading is a label, not a paragraph.
    """
    stripped = line.strip()
    # Markdown heading and emphasis markers, plus a trailing colon.
    stripped = re.sub(r"^\s*#{1,6}\s*", "", stripped)
    stripped = stripped.strip("*_ \t").strip().rstrip(":").strip("*_ \t").strip()
    if not stripped or len(stripped) > 40:
        return None
    # Digits mean a date or a bullet, not a section label.
    if any(character.isdigit() for character in stripped):
        return None

    lowered = stripped.lower()

    # Checked ahead of the general vocabulary: "Languages & Frameworks" has to
    # read as a skills sub-heading before it can be read as the Languages
    # section. A bare "Languages"/"Langues" carries no qualifier and so is
    # unaffected; French "Langages" is programming-specific on its own.
    if _PROGRAMMING_LANGUAGE_WORD.search(lowered):
        return "skills"
    if _LANGUAGE_WORD.search(lowered) and _SKILLS_SUBHEADING.search(lowered):
        return "skills"

    for name, needles in SECTION_PATTERNS.items():
        for needle in needles:
            if re.search(rf"(?<![\w']){re.escape(needle)}s?(?![\w'])", lowered):
                return name
    return None


def _split_heading(line: str) -> tuple[str | None, str]:
    """Split "Label: content" into its section and the content left on the line.

    A heading is normally alone on its line, but three common shapes put real
    content beside it, and treating the whole line as a heading silently ate
    that content:

        Languages & Frameworks: Python, Django    <- a wrapped skills group
        Profile: Engineering student...           <- a one-line CV
        Skills: Python, SQL

    The first is not hypothetical: the `classic` template prints skills as
    "Category: items" and wraps mid-list, so extracting a CV this service had
    itself produced lost the first technology of every group — "Python" simply
    vanished from a round trip.

    Returns the section name and whatever followed the colon, so the caller can
    open the section *and* keep the text. Only splits when the part before the
    colon reads as a heading on its own; "Built a platform: it worked" is
    ordinary prose and stays whole.
    """
    # The colon is checked *first*. "Languages & Frameworks: Python," is short
    # enough and digit-free enough to satisfy the whole-line heading test on
    # its own, so asking that question first would classify it as a bare
    # heading and drop "Python" — which is exactly the bug.
    label, separator, rest = line.partition(":")
    if separator and rest.strip():
        labelled = _looks_like_heading(label)
        if labelled:
            return labelled, rest.strip()
        return None, ""

    return _looks_like_heading(line), ""


# Sections whose contents are self-evident enough to check. If one of these
# holds nothing of the kind its own heading promises, the split is wrong —
# the labels came from somewhere other than the text beneath them.
_LANGUAGE_HINTS = re.compile(
    r"\b(?:english|french|arabic|spanish|german|italian|portuguese|chinese|"
    r"anglais|fran[cç]ais|arabe|espagnol|allemand|"
    r"native|fluent|bilingual|mother\s*tongue|natif|maternelle|courant|"
    r"basic|beginner|intermediate|advanced|notions|d[ée]butant|"
    r"[abc][12])\b",
    re.I,
)
_EDUCATION_HINTS = re.compile(
    r"\b(?:university|universit[ée]|school|[ée]cole|college|institut\w*|"
    r"bachelor|master|licence|diplom\w*|degree|bac\+?\d?|phd|doctorat|"
    r"engineering|ing[ée]nieur|19\d{2}|20\d{2})\b",
    re.I,
)


def _section_is_incoherent(name: str, body: str) -> bool:
    """True when a section plainly does not contain what its heading claims."""
    text = (body or "").strip()
    if not text:
        return False
    if name == "contact":
        # A contact block always carries at least one machine-checkable thing.
        return not (EMAIL_RE.search(text) or PHONE_RE.search(text) or URL_RE.search(text))
    if name == "languages":
        return not _LANGUAGE_HINTS.search(text)
    if name == "education":
        return not _EDUCATION_HINTS.search(text)
    return False


def _best_pdf_text(data: bytes) -> str:
    """Naive drawing-order text, or reading-order text when that reads better.

    `pypdf` emits text in drawing order, which for a two-column CV interleaves
    the sidebar with the main column and hands the section splitter labelled
    sections whose contents belong somewhere else (see `layout.py`). The
    reading-order reconstruction is best-effort and cannot be trusted blindly,
    so both candidates are scored on the same property — headings that
    actually own a body — and the winner is used. A tie keeps the naive text,
    so a file today's code already handles well can only stay the same.
    """
    naive = _text_from_pdf(data)
    try:
        reordered = layout.text_in_reading_order(data)
    except Exception:  # noqa: BLE001 — a reconstruction failure is not fatal
        return naive
    if not reordered.strip():
        return naive

    naive_clean = _clean(naive)
    reordered_clean = _clean(reordered)

    # Reordering may only change the ORDER of the text, never its substance.
    # A reconstruction that drops a section or prints the masthead three times
    # is worse than a scrambled-but-complete read, and the heading score alone
    # would not notice either. Word counts are compared rather than exact
    # strings because line breaks legitimately move when columns are split.
    if not _preserves_content(naive_clean, reordered_clean):
        return naive

    if layout.score_layout(
        reordered_clean.splitlines(), _looks_like_heading
    ) > layout.score_layout(naive_clean.splitlines(), _looks_like_heading):
        return reordered
    return naive


# How far the reordered text's word count may drift from the original before
# it is rejected: a little slack absorbs whitespace rejoining at column edges,
# while real duplication or a lost column moves it far past this.
_CONTENT_DRIFT_TOLERANCE = 0.12


def _preserves_content(original: str, candidate: str) -> bool:
    original_words = original.split()
    if not original_words:
        return False
    candidate_words = candidate.split()
    drift = abs(len(candidate_words) - len(original_words)) / len(original_words)
    if drift > _CONTENT_DRIFT_TOLERANCE:
        return False
    # Cheap check that it is the same text and not merely the same length.
    return len(set(candidate_words) & set(original_words)) >= len(set(original_words)) * 0.9


def extract_cv(data: bytes, filename: str) -> dict:
    """Turn an uploaded CV into labelled, capped sections.

    Returns a dict with `sections` (only those found), `contact` candidates,
    and `notes` describing what the heuristics did — the model is told to treat
    all of it as unconfirmed and check with the user.
    """
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        raw = _best_pdf_text(data)
    elif name.endswith((".docx", ".doc")):
        raw = _text_from_docx(data)
    elif name.endswith((".txt", ".md")):
        raw = data.decode("utf-8", errors="replace")
    else:
        raise ExtractionError(
            "Unsupported file type. Upload a PDF, DOCX, TXT or MD file."
        )

    text = _clean(raw)
    if not text.strip():
        raise ExtractionError(
            "That file has no extractable text — it is probably a scan or an "
            "image-only PDF. Type the details instead and I'll build from those."
        )

    lines = text.splitlines()

    # --- contact candidates, gathered from the whole document ---------------
    emails = EMAIL_RE.findall(text)
    urls = URL_RE.findall(text)
    # Phone matching is noisy: dates and ID numbers match too. Require enough
    # digits to be a real number, and prefer a line that looks like contact.
    phones = [p.strip() for p in PHONE_RE.findall(text) if len(re.sub(r"\D", "", p)) >= 8]

    # The name is nearly always the first substantial line that is not contact.
    #
    # A two-column sidebar layout breaks this assumption: pypdf's extractor
    # follows the PDF's internal drawing order, not the visual reading order,
    # and a design where every section LABEL is drawn before any section BODY
    # (seen on a real upload) puts "Contact / Language / Skills / Experience /
    # About Me" all inside the first six lines, ahead of the actual name. A
    # heading-shaped candidate — including a fragment like "About M e" that
    # still matches the "about" needle — is never the name, so it is skipped
    # rather than accepted as one; the model still has to find the real name
    # itself, but at least is not handed a confidently wrong one.
    estimated_name = ""
    for line in lines[:6]:
        if EMAIL_RE.search(line) or URL_RE.search(line):
            continue
        if len(re.sub(r"\D", "", line)) >= 6:
            continue
        if _looks_like_heading(line):
            continue
        if 2 <= len(line.split()) <= 5 and len(line) <= 50:
            estimated_name = line
            break

    # --- split into sections by heading ------------------------------------
    sections: dict[str, list[str]] = {}
    current: str | None = None
    preamble: list[str] = []
    for line in lines:
        heading, inline = _split_heading(line)
        if heading:
            current = heading
            sections.setdefault(current, [])
            # "Skills: Python, SQL" opens the section *and* contributes its
            # first line, rather than the label eating the content.
            if inline:
                sections[current].append(inline)
            continue
        (sections[current] if current else preamble).append(line)

    trimmed = {
        key: "\n".join(value).strip()[:SECTION_CAP]
        for key, value in sections.items()
        if "\n".join(value).strip()
    }
    notes: list[str] = []
    if estimated_name:
        notes.append(f"Name is probably '{estimated_name}' (first line).")

    # A heading that owns text belonging to a different part of the page is
    # the worst possible outcome: it looks parsed, so the model trusts it, and
    # what it then "reads" is somebody else's section. When even one section
    # plainly contradicts its own label the split cannot be trusted at all, so
    # the labels are dropped and the whole document is handed over unsplit —
    # the same treatment as a CV with no recognisable headings, which the
    # model already knows how to handle.
    incoherent = [key for key, body in trimmed.items() if _section_is_incoherent(key, body)]
    if incoherent:
        notes.append(
            f"The heading split looked wrong ({', '.join(sorted(incoherent))} did not "
            "contain what that heading promises), so it was discarded and the text "
            "is unsplit under 'full_text'. This usually means a two-column layout "
            "was read out of order. Read it yourself and map it to the CV fields."
        )
        trimmed = {"full_text": text[:TOTAL_CAP]}

    if incoherent:
        # Already handled above, and deliberately not described as "recognised
        # sections" below — the whole point is that nothing was recognised
        # reliably.
        pass
    elif not trimmed:
        # No heading matched, so every line ended up in `preamble` — the whole
        # CV. It goes through the fallback bucket, which has the larger cap;
        # treating it as a masthead would truncate the document to SECTION_CAP.
        notes.append(
            "No section headings were recognised, so the text is unsplit under "
            "'full_text'. Read it and map it to the CV fields yourself."
        )
        trimmed["full_text"] = text[:TOTAL_CAP]
    else:
        # Everything above the first heading. This was collected and then
        # dropped, which quietly lost the masthead — the name and, right under
        # it, the professional title ("MARKETING MANAGER", "AI & Data
        # Engineering"). That title has nowhere else to come from, so every
        # rebuilt CV came out with a bare name.
        masthead = "\n".join(preamble).strip()
        if masthead:
            # First, because it is the top of the page and reads that way.
            trimmed = {"header": masthead[:SECTION_CAP], **trimmed}
            notes.append(
                "'header' is the text above the first heading: it usually holds "
                "the name and the professional title that belongs in `headline`."
            )
        notes.append(f"Recognised sections: {', '.join(sorted(trimmed))}.")

    contact_bits: list[str] = []
    if emails:
        contact_bits.append(emails[0])
    if phones:
        contact_bits.append(phones[0].strip())
    for url in urls[:2]:
        contact_bits.append(url.rstrip(".,);"))

    result = {
        "estimated_name": estimated_name,
        "contact_candidates": contact_bits,
        "sections": trimmed,
        "notes": notes,
        "characters": len(text),
        # Text came out, but not in an order that can be trusted. The caller
        # uses this to escalate to vision even though the grade is not FAILED.
        "layout_unreliable": bool(incoherent),
    }
    result["assessment"] = assess(result).as_dict()
    return result


def extract_everything(data: bytes, filename: str) -> dict:
    """Full first pass over an upload: text, quality grade, and the portrait.

    This is the deterministic tier of the cascade — no model is called here.
    The caller reads `assessment.grade` to decide whether that was enough
    (`good` / `partial`) or whether the file needs the vision fallback
    (`failed`). See `cv/quality.py` for why that split is worth making.

    A photo is a bonus, never a requirement: failing to find one is silent,
    and failing to *read* one must not fail the upload.
    """
    is_pdf = (filename or "").lower().endswith(".pdf")

    try:
        result = extract_cv(data, filename)
    except ExtractionError:
        if not is_pdf:
            raise
        # No text at all. Not fatal on its own — a scanned CV still has pixels,
        # and the vision tier can read them — so report it as a failed grade
        # rather than an error and let the caller decide.
        result = {
            "estimated_name": "",
            "contact_candidates": [],
            "sections": {},
            "notes": ["No text could be extracted from this PDF."],
            "characters": 0,
            "assessment": Assessment(
                Grade.FAILED,
                ["no extractable text at all — the file is an image or a scan"],
            ).as_dict(),
        }

    if is_pdf:
        result["photo"] = extract_portrait(data)
    elif (filename or "").lower().endswith(".docx"):
        result["photo"] = extract_portrait_from_docx(data)
    else:
        result["photo"] = None
    # Vision input, fetched only when the text tier could not be trusted —
    # decoding a full page raster costs real memory, so it is never done
    # speculatively.
    #
    # Two different failures qualify, not one. `failed` means no usable text
    # came out (a scan). `layout_unreliable` means plenty of text came out in
    # the wrong order — a two-column design whose sidebar interleaves with the
    # main column, which produced confidently-labelled sections holding other
    # sections' content. The second case is the more dangerous of the two,
    # because it looks like a successful parse.
    assessment = result.get("assessment") or {}
    needs_vision = assessment.get("grade") == "failed" or result.get("layout_unreliable")
    if is_pdf and needs_vision:
        # An embedded image is the cheaper read when the PDF is a scan; a text
        # PDF has none, so the page is drawn instead.
        result["page_image"] = extract_page_image(data) or render_pdf_page(data)
    else:
        result["page_image"] = None
    return result
