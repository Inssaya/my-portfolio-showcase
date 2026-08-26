"""The extraction cascade: grade first, spend on vision only if it must.

The routing decision is the whole economic argument — deterministic parsing is
free, a vision call is not — so each branch is pinned here, along with the
photo extraction that feeds the rebuilt CV.
"""
from __future__ import annotations

import io

import pytest

from app.cv.builder import build_resume
from app.cv.extract import extract_everything
from app.cv.photo import extract_page_image, extract_portrait
from app.cv.quality import Grade, assess

READABLE_CV = b"""\
Yassine Sinif
yassinsinif4@gmail.com | +212 6 23 84 25 35

Profile
Final-year engineering student in AI and Data Science.

Experience
AI Data Engineer Intern, Aptiv, 2026
- Built a maintenance KPI platform.

Education
Engineering Degree, EMSI Casablanca

Skills
Python, TypeScript, Kafka
"""


# ------------------------------------------------------------------ grading

def test_a_well_formed_cv_grades_good() -> None:
    """The cheap path must be the common one, or the economics do not work."""
    result = extract_everything(READABLE_CV, "cv.txt")

    assert result["assessment"]["grade"] == Grade.GOOD.value
    assert result["assessment"]["has_contact"] is True


def test_structureless_text_grades_partial_not_failed() -> None:
    """Readable but unstructured is the model's job, not vision's."""
    body = ("I worked at Acme for three years doing marketing and strategy. " * 12).encode()
    result = extract_everything(body, "cv.txt")

    assert result["assessment"]["grade"] == Grade.PARTIAL.value


def test_short_cv_with_structure_is_not_rejected() -> None:
    """Regression: a length-only rule graded a short but perfectly readable CV
    as a scan, which refused the upload outright."""
    tiny = b"Jane Doe\njane@example.com\n\nExperience\nBuilt things at Acme.\n"
    result = extract_everything(tiny, "cv.txt")

    assert result["assessment"]["grade"] != Grade.FAILED.value


def test_empty_extraction_grades_failed() -> None:
    assessment = assess({"sections": {}, "characters": 0, "contact_candidates": []})

    assert assessment.grade is Grade.FAILED
    assert assessment.needs_vision is True


@pytest.mark.parametrize(
    "noise",
    [
        # A font with no ToUnicode map: pypdf emits glyph ids like this.
        " ".join(["(cid:11)", "(cid:24)", "(cid:3)", "(cid:87)"] * 100),
        # A broken encoding maps everything into the private use area.
        " ".join(["", "", ""] * 120),
    ],
    ids=["cid-refs", "private-use-area"],
)
def test_gibberish_grades_failed(noise: str) -> None:
    """A broken embedded font yields plenty of characters and no words — the
    case a raw character count waves straight through.

    Note what this does and does not catch: it tests whether tokens are
    *shaped* like words, not whether they are real words. Letter-salad
    ("xkqz vbnm") passes, and deliberately so — ruling that out needs a
    dictionary per language, while the actual failure mode in extracted PDFs is
    symbol garbage, which this catches.
    """
    assessment = assess(
        {"sections": {"full_text": noise}, "characters": len(noise), "contact_candidates": []}
    )

    assert assessment.grade is Grade.FAILED
    assert "gibberish" in " ".join(assessment.reasons)


def test_good_grade_needs_no_vision() -> None:
    assert assess(
        {
            "sections": {"experience": "Worked at Acme doing real things.", "skills": "Python"},
            "characters": 900,
            "contact_candidates": ["a@b.com"],
        }
    ).needs_vision is False



# The two round-trip tests below used to point `photo=` at
# "C:/Users/yassi/portfolio/cv/photo.png" — a path that exists only on the
# machine the test was written on. Everywhere else `build_resume` silently
# skipped the photo (a bad portrait must never lose the CV, by design), so the
# tests failed on the assertion that a photo came back out, and had been
# failing on every other host and in CI ever since. Generating the portrait
# here keeps them testing the round trip they claim to test, on any machine.
@pytest.fixture
def portrait_path(tmp_path):
    from PIL import Image, ImageDraw

    image = Image.new("RGB", (420, 420), (238, 226, 214))
    draw = ImageDraw.Draw(image)
    draw.ellipse([110, 70, 310, 270], fill=(176, 137, 116))
    draw.rectangle([80, 255, 340, 420], fill=(58, 72, 96))
    path = tmp_path / "portrait.png"
    image.save(path, format="PNG")
    return str(path)


# -------------------------------------------------------------------- photo

def test_portrait_is_extracted_from_a_generated_cv(portrait_path) -> None:
    """Round-trip: a CV built with a photo must give that photo back."""
    pdf_bytes, _ = build_resume(
        full_name="Yassine Sinif",
        photo=portrait_path,
        skills="Languages: Python",
    )

    portrait = extract_portrait(pdf_bytes)

    assert portrait is not None, "photo went in but did not come out"
    assert portrait[:8] == b"\x89PNG\r\n\x1a\n"

    from PIL import Image

    image = Image.open(io.BytesIO(portrait))
    assert min(image.size) >= 90
    # Roughly square, as a headshot is.
    assert 0.5 <= image.size[0] / image.size[1] <= 1.6


def test_no_photo_is_not_an_error() -> None:
    """Most CVs have no portrait; that is an ordinary outcome, not a failure."""
    pdf_bytes, _ = build_resume(full_name="Yassine Sinif", skills="Languages: Python")

    assert extract_portrait(pdf_bytes) is None


def test_photo_extraction_survives_a_broken_file() -> None:
    """A malformed upload must cost the visitor their photo, not the process."""
    assert extract_portrait(b"%PDF-1.4 not really a pdf") is None
    assert extract_page_image(b"") is None


def test_page_image_ignores_a_small_portrait(portrait_path) -> None:
    """extract_page_image feeds vision. A headshot tells vision nothing, so it
    must not trigger a paid call — this is why it is separate from
    extract_portrait, which deliberately rejects page-sized images.

    Previously the photo path did not resolve, so nothing was embedded and
    this asserted None against a CV with no image in it at all — passing
    without exercising the rejection it exists to pin.
    """
    pdf_bytes, _ = build_resume(
        full_name="Yassine Sinif",
        photo=portrait_path,
        skills="Languages: Python",
    )
    # The portrait really is in there; extract_page_image must still decline it.
    assert extract_portrait(pdf_bytes) is not None

    assert extract_page_image(pdf_bytes) is None


def test_page_image_is_only_fetched_when_vision_is_needed() -> None:
    """Decoding a full-page raster costs real memory; a readable CV must not
    pay it."""
    result = extract_everything(READABLE_CV, "cv.txt")

    assert result["page_image"] is None


# ------------------------------------------------------- photo into the CV

def test_extracted_photo_can_be_rebuilt_into_a_new_cv(portrait_path) -> None:
    """The point of extracting it: a rebuild keeps the visitor's face."""
    import os
    import tempfile

    original, _ = build_resume(
        full_name="Yassine Sinif",
        photo=portrait_path,
        skills="Languages: Python",
    )
    portrait = extract_portrait(original)
    assert portrait is not None

    handle = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    try:
        handle.write(portrait)
        handle.close()
        rebuilt, pages = build_resume(
            full_name="Yassine Sinif", photo=handle.name, skills="Languages: Python"
        )
    finally:
        os.unlink(handle.name)

    assert rebuilt.startswith(b"%PDF-")
    assert pages == 1
    assert extract_portrait(rebuilt) is not None, "photo lost on rebuild"


def test_docx_photo_is_extracted(portrait_path) -> None:
    """Word is what most people write a CV in; without this every .docx upload
    silently loses its photo on rebuild.

    Built here rather than globbed out of the author's Downloads folder, which
    is what this did before — so it skipped on every other machine and in CI,
    leaving the .docx photo path with no coverage at all.
    """
    import docx

    from app.cv.photo import extract_portrait_from_docx

    document = docx.Document()
    document.add_paragraph("Yassine Sinif")
    document.add_picture(portrait_path)
    buffer = io.BytesIO()
    document.save(buffer)

    portrait = extract_portrait_from_docx(buffer.getvalue())

    assert portrait is not None
    assert portrait[:8] == b"\x89PNG\r\n\x1a\n"


def test_docx_photo_extraction_survives_a_non_zip() -> None:
    from app.cv.photo import extract_portrait_from_docx

    assert extract_portrait_from_docx(b"not a zip at all") is None


# --------------------------------------------------- masthead and contact

def test_masthead_above_the_first_heading_is_kept() -> None:
    """Regression: the preamble was collected and then dropped, losing the name
    and the professional title under it — so every rebuild had a bare name."""
    cv = b"""\
AHMDD SAAH
MARKETING MANAGER

PROFILE
Ten years leading brand strategy.

EXPERIENCE
Borcelle Studio, Marketing Manager, 2030 - Present
- Ran the campaigns.
"""
    result = extract_everything(cv, "cv.txt")

    assert "header" in result["sections"]
    assert "MARKETING MANAGER" in result["sections"]["header"]
    # It must read first, as it does on the page.
    assert list(result["sections"])[0] == "header"


def test_contact_heading_is_recognised() -> None:
    """Regression: 'contact' was missing from the heading vocabulary, so the
    whole block was discarded — taking the city with it, since a city matches
    no email, phone or URL pattern."""
    cv = b"""\
Jane Doe

CONTACT
Casablanca, Morocco
jane@example.com

EXPERIENCE
Acme, 2024
- Did the work.
"""
    result = extract_everything(cv, "cv.txt")

    assert "contact" in result["sections"]
    assert "Casablanca, Morocco" in result["sections"]["contact"]


def test_header_does_not_inflate_the_quality_grade() -> None:
    """`header` exists even when no heading was recognised, so counting it as a
    section would let an unstructured CV grade GOOD and skip the warning."""
    result = extract_everything(
        b"Jane Doe\nSome prose about a career with an email jane@example.com\n" * 8,
        "cv.txt",
    )

    assert result["assessment"]["grade"] == Grade.PARTIAL.value


def test_headingless_cv_is_not_truncated_into_the_masthead() -> None:
    """With no headings every line lands in the preamble. It must go through
    the full_text bucket, which has the larger cap, not be cut to SECTION_CAP."""
    long_cv = ("Jane Doe worked at Acme doing many notable things. " * 120).encode()
    result = extract_everything(long_cv, "cv.txt")

    assert "full_text" in result["sections"]
    assert len(result["sections"]["full_text"]) > 1200
